import streamlit as st
import pandas as pd
import xml.etree.ElementTree as ET
import pdfplumber

st.set_page_config(page_title="Auditor de presupuestos de obra", layout="wide")
import streamlit as st
import pandas as pd
import xml.etree.ElementTree as ET

st.set_page_config(page_title="Auditor de presupuestos de obra", layout="wide")


# =========================
# PARSER UNIVERSAL BC3
# =========================

def detectar_formato_bc3(text: str) -> str:
    lines = text.splitlines()
    sample = lines[:200]

    # Presto (~V|, ~C|, ~D|, ~T|)
    if any(l.startswith("~V|") for l in sample) and any(l.startswith("~C|") for l in sample):
        return "presto"

    # Arquímedes (C;..., L;..., D;...)
    if any(l.startswith("C;") or l.startswith("L;") for l in sample):
        return "arquimedes"

    # CYPE (|P| en la segunda columna)
    if any("|P|" in l for l in sample):
        return "cype"

    # XML (TCQ, Menfis, etc.)
    if any("<ITEM>" in l or "<BC3" in l for l in sample):
        return "xml"

    return "desconocido"


def parse_presto(text: str) -> pd.DataFrame:
    """
    Adaptado al formato que has pegado:
    - ~C|codigo|unidad|texto|precio...  -> partidas
    - ~T|codigo|texto_largo             -> texto detallado
    - ~D|codigo_capitulo|...           -> descomposición (no la usamos aún)
    """
    partidas = {}
    lines = text.splitlines()

    for line in lines:
        if line.startswith("~C|"):
            parts = line.split("|")
            # Ejemplos:
            # ~C|001#||DEMOLICIONES|5463.96|...
            # ~C|002001|Ud|Verificación/reparación instalación eléctrica|909.28|...
            if len(parts) >= 5:
                codigo = parts[1].strip()
                unidad = parts[2].strip()
                texto = parts[3].strip()
                try:
                    precio = float(parts[4].replace(",", "."))
                except ValueError:
                    precio = 0.0

                # Para partidas tipo Ud, tomamos cantidad=1 y importe=precio (MVP)
                cantidad = 1.0
                importe = precio

                partidas[codigo] = {
                    "codigo": codigo,
                    "texto": texto,
                    "unidad": unidad,
                    "cantidad": cantidad,
                    "precio": precio,
                    "importe": importe,
                    "descomp_materiales": [],
                    "descomp_mano_obra": [],
                    "descomp_maquinaria": [],
                }

        elif line.startswith("~T|"):
            parts = line.split("|")
            # ~T|002001|Texto largo...
            if len(parts) >= 3:
                codigo = parts[1].strip()
                texto_largo = parts[2].strip()
                if codigo in partidas:
                    # concatenamos texto corto + texto largo
                    base = partidas[codigo]["texto"]
                    if base:
                        partidas[codigo]["texto"] = base + " " + texto_largo
                    else:
                        partidas[codigo]["texto"] = texto_largo

        # ~D|... lo podríamos usar para descomposición, pero de momento lo dejamos fuera
        # porque ya tienes bastante con detectar las partidas.

    df = pd.DataFrame(list(partidas.values()))
    return df


def parse_arquimedes(text: str) -> pd.DataFrame:
    partidas = []

    for line in text.splitlines():
        parts = line.split(";")
        if len(parts) < 6:
            continue

        tipo = parts[0].strip().upper()

        if tipo == "L":  # Partida
            try:
                cantidad = float(parts[4].replace(",", "."))
            except ValueError:
                cantidad = 0.0
            try:
                precio = float(parts[5].replace(",", "."))
            except ValueError:
                precio = 0.0

            partidas.append(
                {
                    "codigo": parts[1].strip(),
                    "texto": parts[2].strip(),
                    "unidad": parts[3].strip(),
                    "cantidad": cantidad,
                    "precio": precio,
                    "importe": cantidad * precio,
                    "descomp_materiales": [],
                    "descomp_mano_obra": [],
                    "descomp_maquinaria": [],
                }
            )

    return pd.DataFrame(partidas)


def parse_cype(text: str) -> pd.DataFrame:
    partidas = []

    for line in text.splitlines():
        parts = line.split("|")
        if len(parts) < 8:
            continue

        tipo = parts[1].strip().upper()
        if tipo == "P":  # Partida
            try:
                cantidad = float(parts[5].replace(",", "."))
            except ValueError:
                cantidad = 0.0
            try:
                precio = float(parts[6].replace(",", "."))
            except ValueError:
                precio = 0.0
            try:
                importe = float(parts[7].replace(",", "."))
            except ValueError:
                importe = cantidad * precio

            partidas.append(
                {
                    "codigo": parts[2].strip(),
                    "texto": parts[3].strip(),
                    "unidad": parts[4].strip(),
                    "cantidad": cantidad,
                    "precio": precio,
                    "importe": importe,
                    "descomp_materiales": [],
                    "descomp_mano_obra": [],
                    "descomp_maquinaria": [],
                }
            )

    return pd.DataFrame(partidas)


def parse_xml(text: str) -> pd.DataFrame:
    partidas = []
    try:
        root = ET.fromstring(text)
    except Exception:
        return pd.DataFrame()

    for item in root.findall(".//ITEM"):
        codigo = item.findtext("CODE", "").strip()
        texto = item.findtext("DESC", "").strip()
        unidad = item.findtext("UNIT", "").strip()
        try:
            cantidad = float(item.findtext("QTY", "0").replace(",", "."))
        except ValueError:
            cantidad = 0.0
        try:
            precio = float(item.findtext("PRICE", "0").replace(",", "."))
        except ValueError:
            precio = 0.0
        importe = cantidad * precio

        partidas.append(
            {
                "codigo": codigo,
                "texto": texto,
                "unidad": unidad,
                "cantidad": cantidad,
                "precio": precio,
                "importe": importe,
                "descomp_materiales": [],
                "descomp_mano_obra": [],
                "descomp_maquinaria": [],
            }
        )

    return pd.DataFrame(partidas)


def parse_bc3_auto(file_bytes: bytes) -> pd.DataFrame:
    text = file_bytes.decode("latin-1", errors="ignore")
    fmt = detectar_formato_bc3(text)

    if fmt == "presto":
        return parse_presto(text)
    if fmt == "arquimedes":
        return parse_arquimedes(text)
    if fmt == "cype":
        return parse_cype(text)
    if fmt == "xml":
        return parse_xml(text)

    raise ValueError("Formato BC3 no reconocido. Ajustar detector o parser.")


# =========================
# MÓDULOS DE ANÁLISIS (13)
# =========================

def es_texto_pobre(texto: str) -> bool:
    if not texto or not isinstance(texto, str):
        return True
    t = texto.strip()
    if len(t) < 10:
        return True
    genericas = ["varios", "completo", "instalación completa", "según proyecto", "trabajos varios"]
    return any(g.lower() in t.lower() for g in genericas)


def detectar_partidas_alzadas(row) -> bool:
    if row["cantidad"] == 0:
        return True
    if str(row["unidad"]).lower() in ["alzada", "global", "servicio", "trabajo"]:
        return True
    if "partida alzada" in str(row["texto"]).lower():
        return True
    return False


def detectar_partidas_incompletas(row) -> bool:
    texto = str(row["texto"]).lower()
    mo = row.get("descomp_mano_obra", [])
    if any(k in texto for k in ["instalación", "montaje", "suministro e instalación"]):
        if len(mo) == 0:
            return True
    return False


def detectar_elementos_sobrantes(row) -> bool:
    texto = str(row["texto"]).lower()
    comps = row.get("descomp_mano_obra", []) + row.get("descomp_materiales", [])
    if "tubería" in texto and any("roza" in c.get("descripcion", "").lower() for c in comps):
        return True
    if "cuadro" in texto and any("cable" in c.get("descripcion", "").lower() for c in comps):
        return True
    return False


def detectar_debe_dividirse(row) -> bool:
    texto = str(row["texto"]).lower()
    patrones = [
        ("cuadro", "cable"),
        ("cuadro", "mecanismo"),
        ("tubería", "aislamiento"),
        ("tubería", "soporte"),
        ("equipo", "instalación"),
        ("bie", "señalización"),
    ]
    for a, b in patrones:
        if a in texto and b in texto:
            return True
    return False


def detectar_sin_medicion(row) -> bool:
    return float(row["cantidad"]) == 0.0


def detectar_sin_descomposicion(row) -> bool:
    return (
        len(row.get("descomp_materiales", [])) == 0
        and len(row.get("descomp_mano_obra", [])) == 0
        and len(row.get("descomp_maquinaria", [])) == 0
    )


def detectar_sin_texto(row) -> bool:
    return es_texto_pobre(row["texto"])


def detectar_duplicadas(df: pd.DataFrame) -> pd.Series:
    clave = (
        df["texto"].fillna("").str.lower()
        + "|"
        + df["unidad"].fillna("").str.lower()
        + "|"
        + df["precio"].fillna(0).astype(str)
    )
    return clave.duplicated(keep=False)


def detectar_contradictorias(df: pd.DataFrame) -> pd.Series:
    contrad = pd.Series(False, index=df.index)
    textos = df["texto"].fillna("").str.lower()
    multicapa = textos.str.contains("multicapa")
    cobre = textos.str.contains("cobre")
    tuberia = textos.str.contains("tubería")
    mask_multi = tuberia & multicapa
    mask_cobre = tuberia & cobre
    if mask_multi.any() and mask_cobre.any():
        contrad = mask_multi | mask_cobre
    return contrad


def extraer_sistemas_desde_actuaciones(doc_text: str) -> list:
    sistemas = []
    claves = ["electricidad", "clima", "climatización", "pci", "fontanería", "saneamiento", "ventilación"]
    for c in claves:
        if c in doc_text.lower():
            sistemas.append(c)
    return list(set(sistemas))


def detectar_faltantes_segun_actuaciones(df: pd.DataFrame, doc_text: str):
    sistemas = extraer_sistemas_desde_actuaciones(doc_text)
    textos = df["texto"].fillna("").str.lower()
    faltantes = []
    for s in sistemas:
        if not textos.str.contains(s).any():
            faltantes.append(s)
    serie = pd.Series(False, index=df.index)
    return serie, faltantes


def detectar_precios_bajos(df: pd.DataFrame) -> pd.Series:
    precios = df["precio"].fillna(0)
    sospechoso = precios < 5

    familias = {
        "tubería": df["texto"].str.contains("tubería", case=False, na=False),
        "cable": df["texto"].str.contains("cable", case=False, na=False),
        "detector": df["texto"].str.contains("detector", case=False, na=False),
        "cuadro": df["texto"].str.contains("cuadro", case=False, na=False),
    }
    for fam, mask in familias.items():
        if mask.any():
            mediana = df.loc[mask, "precio"].median()
            sospechoso |= (mask & (df["precio"] < 0.5 * mediana))

    return sospechoso


def detectar_mano_obra_irreal(row) -> bool:
    texto = str(row["texto"]).lower()
    mo = row.get("descomp_mano_obra", [])

    if any(k in texto for k in ["instalación", "montaje", "suministro e instalación"]):
        if len(mo) == 0:
            return True

    for c in mo:
        if c.get("cantidad", 0) < 0.01:
            return True

    return False


def detectar_partidas_criticas_coste(df: pd.DataFrame) -> pd.Series:
    importe = df["importe"].fillna(df["cantidad"] * df["precio"])
    total = float(importe.sum()) if len(df) else 0.0
    crit = pd.Series(False, index=df.index)

    crit |= importe > 5000
    if total > 0:
        crit |= importe > 0.1 * total

    crit |= (df["precio"] < 5) & (df["cantidad"] > 100)

    return crit


# =========================
# INTERFAZ STREAMLIT
# =========================

st.title("Auditor de presupuestos de obra (instalaciones)")

st.markdown(
    """
Analiza un **BC3** y detecta:

- Partidas alzadas  
- Partidas incompletas  
- Partidas con elementos que sobran  
- Partidas que deberían dividirse  
- Partidas sin medición  
- Partidas sin descomposición  
- Partidas sin texto informativo  
- Partidas duplicadas  
- Partidas contradictorias  
- Partidas faltantes según actuaciones  
- Precios bajos  
- Mano de obra inexistente o irreal  
- Partidas críticas por coste  
"""
)

col1, col2 = st.columns(2)

with col1:
    bc3_file = st.file_uploader("Sube el archivo BC3", type=["bc3", "BC3", "txt"])
with col2:
    docx_file = st.file_uploader("Sube el DOCX de actuaciones (opcional)", type=["docx"])

doc_text = ""
faltantes_global = []
if docx_file is not None:
    try:
        import docx as docx_lib

        doc = docx_lib.Document(docx_file)
        doc_text = "\n".join(p.text for p in doc.paragraphs)
    except Exception:
        st.warning("No se ha podido leer el DOCX. Se omite el análisis de actuaciones.")

if bc3_file is None:
    st.info("Sube un BC3 para empezar el análisis.")
    st.stop()

try:
    df = parse_bc3_auto(bc3_file.read())
except ValueError as e:
    st.error(str(e))
    st.stop()

if df.empty:
    st.warning("No se han detectado partidas en el BC3. Revisa el formato o ajusta el parser.")
    st.stop()

for col in ["codigo", "texto", "unidad", "cantidad", "precio", "importe"]:
    if col not in df.columns:
        df[col] = "" if col in ["codigo", "texto", "unidad"] else 0.0

df["alzada"] = df.apply(detectar_partidas_alzadas, axis=1)
df["incompleta"] = df.apply(detectar_partidas_incompletas, axis=1)
df["sobran_elementos"] = df.apply(detectar_elementos_sobrantes, axis=1)
df["debe_dividirse"] = df.apply(detectar_debe_dividirse, axis=1)
df["sin_medicion"] = df.apply(detectar_sin_medicion, axis=1)
df["sin_descomposicion"] = df.apply(detectar_sin_descomposicion, axis=1)
df["sin_texto"] = df.apply(detectar_sin_texto, axis=1)
df["precio_bajo"] = detectar_precios_bajos(df)
df["mo_irreal"] = df.apply(detectar_mano_obra_irreal, axis=1)
df["critica_coste"] = detectar_partidas_criticas_coste(df)

df["duplicada"] = detectar_duplicadas(df)
df["contradictoria"] = detectar_contradictorias(df)

if doc_text:
    serie_faltantes, faltantes_global = detectar_faltantes_segun_actuaciones(df, doc_text)
    df["faltante_segun_actuaciones"] = serie_faltantes
else:
    df["faltante_segun_actuaciones"] = False

with st.sidebar:
    st.header("Resumen de alertas")

    def count(col):
        return int(df[col].sum())

    st.write(f"🔴 Partidas alzadas: **{count('alzada')}**")
    st.write(f"🔴 Incompletas: **{count('incompleta')}**")
    st.write(f"🔴 Con elementos que sobran: **{count('sobran_elementos')}**")
    st.write(f"🔴 Deben dividirse: **{count('debe_dividirse')}**")
    st.write(f"🔴 Sin medición: **{count('sin_medicion')}**")
    st.write(f"🔴 Sin descomposición: **{count('sin_descomposicion')}**")
    st.write(f"🔴 Sin texto informativo: **{count('sin_texto')}**")
    st.write(f"🔴 Duplicadas: **{count('duplicada')}**")
    st.write(f"🔴 Contradictorias: **{count('contradictoria')}**")
    st.write(f"🔴 Precios bajos: **{count('precio_bajo')}**")
    st.write(f"🔴 Mano de obra inexistente/irreal: **{count('mo_irreal')}**")
    st.write(f"🔴 Críticas por coste: **{count('critica_coste')}**")

    if faltantes_global:
        st.markdown("### Sistemas faltantes según actuaciones:")
        for s in faltantes_global:
            st.write(f"- {s}")

st.subheader("Filtros de visualización")

cols_alertas = [
    "alzada",
    "incompleta",
    "sobran_elementos",
    "debe_dividirse",
    "sin_medicion",
    "sin_descomposicion",
    "sin_texto",
    "duplicada",
    "contradictoria",
    "faltante_segun_actuaciones",
    "precio_bajo",
    "mo_irreal",
    "critica_coste",
]

alerta_sel = st.multiselect(
    "Mostrar solo partidas con estas alertas:",
    options=cols_alertas,
    default=[],
)

df_view = df.copy()
if alerta_sel:
    mask = False
    for c in alerta_sel:
        mask |= df_view[c]
    df_view = df_view[mask]

st.subheader("Partidas analizadas")

cols_mostrar = [
    "codigo",
    "texto",
    "unidad",
    "cantidad",
    "precio",
    "importe",
] + cols_alertas

st.dataframe(df_view[cols_mostrar].reset_index(drop=True), use_container_width=True)


# =========================
# PARSER UNIVERSAL BC3
# =========================

def detectar_formato_bc3(text: str) -> str:
    lines = text.splitlines()
    sample = lines[:200]

    if any(l.startswith("~V|") for l in sample) and any(l.startswith("~C|") for l in sample):
        return "presto"

    if any(l.startswith("C;") or l.startswith("L;") for l in sample):
        return "arquimedes"

    if any("|P|" in l for l in sample):
        return "cype"

    if any("<ITEM>" in l or "<BC3" in l for l in sample):
        return "xml"

    return "desconocido"


def parse_presto(text: str) -> pd.DataFrame:
    partidas = {}
    lines = text.splitlines()

    for line in lines:
        if line.startswith("~C|"):
            parts = line.split("|")
            if len(parts) >= 5:
                codigo = parts[1].strip()
                unidad = parts[2].strip()
                texto = parts[3].strip()
                try:
                    precio = float(parts[4].replace(",", "."))
                except ValueError:
                    precio = 0.0

                cantidad = 1.0
                importe = precio

                partidas[codigo] = {
                    "codigo": codigo,
                    "texto": texto,
                    "unidad": unidad,
                    "cantidad": cantidad,
                    "precio": precio,
                    "importe": importe,
                    "descomp_materiales": [],
                    "descomp_mano_obra": [],
                    "descomp_maquinaria": [],
                }

        elif line.startswith("~T|"):
            parts = line.split("|")
            if len(parts) >= 3:
                codigo = parts[1].strip()
                texto_largo = parts[2].strip()
                if codigo in partidas:
                    base = partidas[codigo]["texto"]
                    if base:
                        partidas[codigo]["texto"] = base + " " + texto_largo
                    else:
                        partidas[codigo]["texto"] = texto_largo

    df = pd.DataFrame(list(partidas.values()))
    return df


def parse_arquimedes(text: str) -> pd.DataFrame:
    partidas = []

    for line in text.splitlines():
        parts = line.split(";")
        if len(parts) < 6:
            continue

        tipo = parts[0].strip().upper()

        if tipo == "L":
            try:
                cantidad = float(parts[4].replace(",", "."))
            except ValueError:
                cantidad = 0.0
            try:
                precio = float(parts[5].replace(",", "."))
            except ValueError:
                precio = 0.0

            partidas.append(
                {
                    "codigo": parts[1].strip(),
                    "texto": parts[2].strip(),
                    "unidad": parts[3].strip(),
                    "cantidad": cantidad,
                    "precio": precio,
                    "importe": cantidad * precio,
                    "descomp_materiales": [],
                    "descomp_mano_obra": [],
                    "descomp_maquinaria": [],
                }
            )

    return pd.DataFrame(partidas)


def parse_cype(text: str) -> pd.DataFrame:
    partidas = []

    for line in text.splitlines():
        parts = line.split("|")
        if len(parts) < 8:
            continue

        tipo = parts[1].strip().upper()
        if tipo == "P":
            try:
                cantidad = float(parts[5].replace(",", "."))
            except ValueError:
                cantidad = 0.0
            try:
                precio = float(parts[6].replace(",", "."))
            except ValueError:
                precio = 0.0
            try:
                importe = float(parts[7].replace(",", "."))
            except ValueError:
                importe = cantidad * precio

            partidas.append(
                {
                    "codigo": parts[2].strip(),
                    "texto": parts[3].strip(),
                    "unidad": parts[4].strip(),
                    "cantidad": cantidad,
                    "precio": precio,
                    "importe": importe,
                    "descomp_materiales": [],
                    "descomp_mano_obra": [],
                    "descomp_maquinaria": [],
                }
            )

    return pd.DataFrame(partidas)


def parse_xml(text: str) -> pd.DataFrame:
    partidas = []
    try:
        root = ET.fromstring(text)
    except Exception:
        return pd.DataFrame()

    for item in root.findall(".//ITEM"):
        codigo = item.findtext("CODE", "").strip()
        texto = item.findtext("DESC", "").strip()
        unidad = item.findtext("UNIT", "").strip()
        try:
            cantidad = float(item.findtext("QTY", "0").replace(",", "."))
        except ValueError:
            cantidad = 0.0
        try:
            precio = float(item.findtext("PRICE", "0").replace(",", "."))
        except ValueError:
            precio = 0.0
        importe = cantidad * precio

        partidas.append(
            {
                "codigo": codigo,
                "texto": texto,
                "unidad": unidad,
                "cantidad": cantidad,
                "precio": precio,
                "importe": importe,
                "descomp_materiales": [],
                "descomp_mano_obra": [],
                "descomp_maquinaria": [],
            }
        )

    return pd.DataFrame(partidas)


def parse_bc3_auto(file_bytes: bytes) -> pd.DataFrame:
    text = file_bytes.decode("latin-1", errors="ignore")
    fmt = detectar_formato_bc3(text)

    if fmt == "presto":
        return parse_presto(text)
    if fmt == "arquimedes":
        return parse_arquimedes(text)
    if fmt == "cype":
        return parse_cype(text)
    if fmt == "xml":
        return parse_xml(text)

    raise ValueError("Formato BC3 no reconocido. Ajustar detector o parser.")


# =========================
# MÓDULOS DE ANÁLISIS (13)
# =========================
# (Tu bloque completo aquí, sin tocar nada)


# =========================
# INTERFAZ STREAMLIT
# =========================

st.title("Auditor de presupuestos de obra (instalaciones)")

col1, col2, col3 = st.columns(3)

with col1:
    bc3_file = st.file_uploader("Sube el archivo BC3", type=["bc3", "BC3", "txt"])
with col2:
    docx_file = st.file_uploader("Sube el DOCX de actuaciones (opcional)", type=["docx"])
with col3:
    pdf_file = st.file_uploader("Sube el PDF de actuaciones (opcional)", type=["pdf"])

doc_text = ""
pdf_text = ""
faltantes_global = []

# DOCX
if docx_file is not None:
    try:
        import docx as docx_lib
        doc = docx_lib.Document(docx_file)
        doc_text = "\n".join(p.text for p in doc.paragraphs)
    except:
        st.warning("No se ha podido leer el DOCX.")

# PDF
if pdf_file is not None:
    try:
        with pdfplumber.open(pdf_file) as pdf:
            pdf_text = "\n".join(page.extract_text() or "" for page in pdf.pages)
    except:
        st.warning("No se ha podido leer el PDF.")

texto_actuaciones = doc_text or pdf_text

if bc3_file is None:
    st.info("Sube un BC3 para empezar el análisis.")
    st.stop()

df = parse_bc3_auto(bc3_file.read())

if df.empty:
    st.warning("No se han detectado partidas en el BC3.")
    st.stop()

# =========================
# APLICAR TUS 13 ANÁLISIS
# =========================
# (Aquí van tus análisis tal cual los tienes)


# =========================
# FALTANTES SEGÚN ACTUACIONES (DOCX o PDF)
# =========================

if texto_actuaciones:
    serie_faltantes, faltantes_global = detectar_faltantes_segun_actuaciones(df, texto_actuaciones)
    df["faltante_segun_actuaciones"] = serie_faltantes
else:
    df["faltante_segun_actuaciones"] = False


# =========================
# TABLA FINAL
# =========================

cols_alertas = [
    "alzada","incompleta","sobran_elementos","debe_dividirse","sin_medicion",
    "sin_descomposicion","sin_texto","duplicada","contradictoria",
    "faltante_segun_actuaciones","precio_bajo","mo_irreal","critica_coste"
]

cols_mostrar = [
    "codigo","texto","unidad","cantidad","precio","importe"
] + cols_alertas

st.dataframe(df[cols_mostrar].reset_index(drop=True), use_container_width=True)
